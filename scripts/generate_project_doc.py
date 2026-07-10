"""Generate Solar Path project overview Word document."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Solar_Path_Project_Overview.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_document() -> Document:
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Solar Path")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Home Energy Platform for Bavaria")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(30, 41, 59)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Project overview · Scope · Technical build\n{date.today().strftime('%B %d, %Y')}")

    doc.add_paragraph()
    live = doc.add_paragraph()
    live.alignment = WD_ALIGN_PARAGRAPH.CENTER
    live.add_run("Live application: ").bold = True
    live.add_run("https://solar-path.onrender.com")

    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.add_run("Source repository: ").bold = True
    repo.add_run("https://github.com/SyedZain295/solarpath")

    doc.add_page_break()

    # 1. Executive summary
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Solar Path is a web application that helps homeowners and businesses in Bavaria (Bayern), "
        "Germany plan home energy systems — solar PV, battery storage, heat pumps, and EV charging. "
        "Users complete a guided energy check, receive sizing recommendations backed by real EU solar "
        "radiation data (PVGIS), compare three honest package options, and connect with verified local "
        "installers for quotes."
    )
    doc.add_paragraph(
        "The platform also serves solar suppliers through a directory, lead pipeline, and subscription "
        "checkout, and includes an EV marketplace module for vehicle matching and home-energy bundling. "
        "The project was built as a production-oriented MVP: monolithic Flask backend, server-rendered "
        "UI, REST APIs, automated testing, CI/CD, and cloud deployment on Render."
    )

    # 2. Project scope
    add_heading(doc, "2. Project Scope")
    add_heading(doc, "2.1 In scope", level=2)
    add_bullet(doc, "Multi-step solar calculator with goals, property, usage, and budget inputs")
    add_bullet(doc, "PV sizing and financial modelling (payback, savings, CO₂)")
    add_bullet(doc, "Three-package comparison (cheapest, best value, most resilient)")
    add_bullet(doc, "Installer directory and postcode-based matching (~18,000 Bavaria-focused entries)")
    add_bullet(doc, "Quote comparison and lead submission workflow")
    add_bullet(doc, "Supplier portal, registration, and Stripe subscription checkout")
    add_bullet(doc, "Customer accounts, PDF decision reports, and bilingual UI (English / German)")
    add_bullet(doc, "Energy advisor, quick estimate, surveys, and compatibility checker")
    add_bullet(doc, "EV marketplace hub (find vehicles, home energy check, bundle planning)")
    add_bullet(doc, "Admin dashboard, health monitoring, beta access gate, and email notifications")

    add_heading(doc, "2.2 Out of scope (documented roadmap)", level=2)
    add_bullet(doc, "Full production CRM for installers")
    add_bullet(doc, "Real-time roof AI analysis from drone imagery")
    add_bullet(doc, "Microservices architecture or multi-region HA")
    add_bullet(doc, "Clinical / FHIR healthcare data (not applicable to solar leads)")

    add_heading(doc, "2.3 Target users", level=2)
    add_table(
        doc,
        ["User type", "Primary need"],
        [
            ["Homeowners & renters", "Understand solar options, costs, and savings for their property"],
            ["Businesses & farms", "Commercial / agri-PV sizing and quote-ready briefs"],
            ["Solar installers", "Receive qualified leads and manage listings"],
            ["EV buyers", "Match vehicles with home charging and solar bundle plans"],
            ["Reviewers / coaches", "Demo the live app without local setup"],
        ],
    )

    # 3. Key features
    add_heading(doc, "3. Key Features")
    add_table(
        doc,
        ["Feature", "Route", "Description"],
        [
            ["Home landing", "/", "Marketing page with process overview and supplier pricing"],
            ["Energy calculator", "/calculator", "8-step wizard: situation, goals, home, location, usage, property, profile, budget"],
            ["Results", "/results", "kWp sizing, packages, savings, matched suppliers, PDF export"],
            ["Demo", "/demo", "Fixed München sample — no calculator required"],
            ["Quick estimate", "/estimate", "60-second entry that pre-fills the full calculator"],
            ["Installers", "/suppliers", "Directory, map, and supplier registration"],
            ["Compare quotes", "/compare-quotes", "Side-by-side installer quote comparison"],
            ["Compatibility", "/compatibility", "Panel / inverter / battery compatibility check"],
            ["Energy advisor", "/energy-advisor", "Guided advice for renters and apartments"],
            ["EV hub", "/ev", "Vehicle finder, listings, home energy, and bundle planner"],
            ["Admin", "/admin", "Platform statistics and management (token-protected)"],
        ],
    )

    # 4. Technology stack
    add_heading(doc, "4. Technology Stack")
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ["Backend", "Python 3.12, Flask 3, Gunicorn, SQLAlchemy"],
            ["Frontend", "Jinja2 templates, vanilla JavaScript, custom CSS design system"],
            ["Database", "SQLite (local dev), PostgreSQL via Neon (production on Render)"],
            ["Solar data", "PVGIS EU JRC API, Global Solar Atlas validation"],
            ["Geocoding", "Nominatim, OpenPLZ, Open-Meteo, cached coordinates"],
            ["Payments", "Stripe (supplier subscriptions)"],
            ["Email", "SMTP (Gmail) for lead and report notifications"],
            ["PDF reports", "ReportLab"],
            ["Testing", "pytest (149 tests), Hypothesis property tests, coverage gates"],
            ["CI/CD", "GitHub Actions (Ubuntu + Windows), pre-commit hooks, CodeQL"],
            ["Deployment", "Docker, Render (render.yaml), GitHub Codespaces for dev"],
        ],
    )

    # 5. How the app was built
    add_heading(doc, "5. How the Application Was Built")
    add_heading(doc, "5.1 Architecture pattern", level=2)
    doc.add_paragraph(
        "Solar Path follows a monolithic Flask architecture: a single application (app.py, ~2,785 lines) "
        "handles page routes, REST JSON APIs, sessions, and orchestration. Domain logic is split into "
        "focused Python modules rather than microservices, which keeps deployment simple while remaining "
        "maintainable for an MVP."
    )

    add_heading(doc, "5.2 Core modules", level=2)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["app.py", "HTTP routing (~100 routes), API endpoints, session handling"],
            ["solar_engine.py", "Orchestrates full recommendation from user inputs"],
            ["decision_engine.py", "Maps user goals to technology recommendations"],
            ["financial_model.py", "Payback, savings, tariffs, Bundesnetzagentur feed-in data"],
            ["pvgis_client.py", "PVGIS yield API with retry, caching, and fallbacks"],
            ["product_catalog.py", "Panels, inverters, batteries, compatibility rules"],
            ["platform_features.py", "Supplier matching, readiness scores, lead routing"],
            ["lead_qualification.py", "Scores and tiers incoming leads for suppliers"],
            ["pdf_report.py", "Bilingual PDF decision report generation"],
            ["database.py", "SQLAlchemy models and persistence layer"],
            ["email_service.py", "SMTP notifications for leads and reports"],
            ["i18n.py / i18n_ev_marketplace.py", "English and German translations"],
            ["beta_access.py", "Optional beta gate and invite tokens"],
            ["compliance.py / security_events.py", "PII handling, audit logging, anomaly detection"],
        ],
    )

    add_heading(doc, "5.3 Calculator data flow", level=2)
    doc.add_paragraph(
        "1. User completes the multi-step calculator in the browser (calculator.js).\n"
        "2. Form data is sent as JSON to POST /api/calculate.\n"
        "3. The server geocodes the postcode if needed and fetches specific solar yield from PVGIS.\n"
        "4. solar_engine.generate_recommendation() combines yield, usage, goals, and roof data.\n"
        "5. financial_model computes payback, annual savings, and three package options.\n"
        "6. platform_features ranks matching installers by distance and quality signals.\n"
        "7. JSON response is stored client-side and displayed on /results.\n"
        "8. User can download a PDF report or submit a quote request to suppliers."
    )

    add_heading(doc, "5.4 Frontend approach", level=2)
    doc.add_paragraph(
        "The UI uses server-rendered Jinja2 templates extended from a shared base layout. "
        "Styling is organised into layered CSS files (style.css, pages.css, aesthetic.css, "
        "design-refinement.css, app-experience.css) with a teal/amber design system. "
        "JavaScript is vanilla ES6+ with no React/Vue framework — each page loads only the "
        "scripts it needs (calculator.js, results.js, survey.js, etc.). Internationalisation "
        "is injected via window.APP_TRANSLATIONS for client-side strings."
    )

    add_heading(doc, "5.5 Data storage", level=2)
    doc.add_paragraph(
        "Installer directory data lives in JSON files (data/suppliers.json, ~18k records from PVR "
        "and OpenStreetMap imports). Product catalog, surveys, quotes, and demo data are also JSON-backed. "
        "Transactional data (customers, subscriptions, EV dealer records) uses SQLite locally and "
        "PostgreSQL in production. Geocoding results are cached in data/geo_cache.json and "
        "data/city_coords.json to reduce external API calls."
    )

    # 6. Development process
    add_heading(doc, "6. Development Process")
    add_bullet(doc, " — Version control with Git; scoped commits and pre-commit hooks for hygiene", "Git workflow")
    add_bullet(doc, " — SETUP_DEV.bat / setup.sh for one-command local setup", "Local dev")
    add_bullet(doc, " — GitHub Codespaces with auto-start for cloud development", "Codespaces")
    add_bullet(doc, " — 149 automated tests run on every push (Ubuntu + Windows matrix)", "Testing")
    add_bullet(doc, " — pip-audit, CodeQL SAST, coverage gates in CI", "Security checks")
    add_bullet(doc, " — render.yaml Infrastructure-as-Code for production deploy", "Deployment")

    add_heading(doc, "6.1 Codebase scale (approximate)", level=2)
    add_table(
        doc,
        ["Component", "Lines of code"],
        [
            ["Python (backend modules)", "~16,200"],
            ["HTML templates", "~4,000"],
            ["CSS", "~11,300"],
            ["JavaScript", "~6,500"],
            ["Automated test files", "25 modules, 149 tests"],
            ["Total production code", "~38,000 lines"],
        ],
    )

    # 7. Deployment
    add_heading(doc, "7. Deployment & Operations")
    doc.add_paragraph(
        "Production runs on Render (free tier web service) with PostgreSQL from Neon. "
        "Gunicorn serves the Flask app with 2 workers. Health checks hit /health which reports "
        "database status, email (SMTP) readiness, and beta gate state. Docker Compose files "
        "support local and Codespaces environments. Environment variables (SECRET_KEY, ADMIN_TOKEN, "
        "DATABASE_URL, SMTP credentials) are never committed to the repository."
    )

    # 8. Security
    add_heading(doc, "8. Security & Compliance")
    add_bullet(doc, "Environment-based secrets with production validation in config.py")
    add_bullet(doc, "bcrypt password hashing for customers, suppliers, and EV dealers")
    add_bullet(doc, "Security headers and Content-Security-Policy in production (http_middleware.py)")
    add_bullet(doc, "Auth anomaly lockout and structured audit logging")
    add_bullet(doc, "PII classification and redaction in logs (compliance.py)")
    add_bullet(doc, "Optional beta access gate with invite tokens for controlled rollouts")

    # 9. Conclusion
    add_heading(doc, "9. Summary")
    doc.add_paragraph(
        "Solar Path is a full-stack home energy platform built with modern Python web practices, "
        "real EU solar data, and a user-centred multi-step calculator experience. It demonstrates "
        "end-to-end product development: domain modelling, external API integration, bilingual UX, "
        "supplier marketplace mechanics, automated testing, and cloud deployment — all within a "
        "single maintainable codebase scoped for the Bavarian solar market."
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
